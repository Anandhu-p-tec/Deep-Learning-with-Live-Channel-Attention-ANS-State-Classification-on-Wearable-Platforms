"""
ANS Project Report - NIET B.E. Final Year Format (Corrected & Complete)
Includes correct technical specifications, formatting, images and references
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import glob

class NIETReportFinal:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.page_num = 0
        self.find_images()
        
    def setup_document(self):
        """Setup NIET standard formatting"""
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
    def find_images(self):
        """Find all project images"""
        self.images = {}
        img_dir = r'c:\Users\ranan\Desktop\ANS\docs\extracted_images'
        if os.path.exists(img_dir):
            for img in glob.glob(os.path.join(img_dir, '*.png')) + glob.glob(os.path.join(img_dir, '*.jpg')):
                name = os.path.basename(img)
                if 'ppt_slide' in name:
                    self.images[name] = img
    
    def add_centered_text(self, text, font_size=14, bold=False, italic=False, space_before=6, space_after=6):
        """Add centered paragraph"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        return p
    
    def add_body_paragraph(self, text, font_size=14, justify=True, space_after=6):
        """Add justified body text"""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
        return p
    
    def add_chapter_heading(self, num, title):
        """Add chapter heading - 16pt bold centered"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CHAPTER {num}: {title.upper()}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.bold = True
        return p
    
    def add_section_heading(self, num, title):
        """Add section heading - 14pt bold left aligned"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{num}. {title}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        return p
    
    def add_image(self, img_path, width=Inches(4.5), caption=None):
        """Add image with caption"""
        try:
            if os.path.exists(img_path):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=width)
                
                if caption:
                    cap_p = self.doc.add_paragraph(caption)
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_before = Pt(3)
                    cap_p.paragraph_format.space_after = Pt(9)
                    for run in cap_p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.italic = True
        except Exception as e:
            print(f"Could not add image {img_path}: {e}")
    
    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()
    
    def create_cover_page(self):
        """NIET Cover Page Format"""
        self.add_centered_text("ON-DEVICE DEEP LEARNING WITH LIVE CHANNEL\nATTENTION FOR REAL-TIME ANS STATE\nCLASSIFICATION ON WEARABLE PLATFORMS", 
                             font_size=18, bold=True, space_before=24, space_after=24)
        
        self.add_centered_text("A PROJECT REPORT", font_size=14)
        self.doc.add_paragraph()
        
        self.add_centered_text("Submitted by", font_size=14, italic=True)
        self.add_centered_text("Anandhu P.", font_size=16, bold=True)
        self.add_centered_text("Ajay B. S.", font_size=16, bold=True) 
        self.add_centered_text("Arun C.", font_size=16, bold=True)
        
        self.doc.add_paragraph()
        self.add_centered_text("in partial fulfillment for the award of the degree of", 
                             font_size=14, italic=True)
        self.add_centered_text("BACHELOR OF ENGINEERING", font_size=16, bold=True, space_before=6, space_after=6)
        self.add_centered_text("in", font_size=14, italic=True)
        self.add_centered_text("ARTIFICIAL INTELLIGENCE AND DATA SCIENCE", font_size=16, bold=True, space_before=6, space_after=24)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text("NEHRU INSTITUTE OF ENGINEERING AND TECHNOLOGY,", font_size=16, bold=True)
        self.add_centered_text("(Autonomous)", font_size=12)
        self.add_centered_text("Affiliated to Anna University, Chennai", font_size=11)
        self.add_centered_text("Nehru Gardens, Thirumalayampalayam, Coimbatore - 641 105", font_size=11)
        
        self.doc.add_paragraph()
        self.add_centered_text("APRIL 2026", font_size=14, bold=True, space_before=12)
        
        self.add_page_break()
    
    def create_certificate(self):
        """Bonafide Certificate - NIET Format"""
        self.add_centered_text("BONAFIDE CERTIFICATE", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        cert_text = """Certified that this project report "ON-DEVICE DEEP LEARNING WITH LIVE CHANNEL ATTENTION FOR REAL-TIME ANS STATE CLASSIFICATION ON WEARABLE PLATFORMS" is the bonafide work of "Anandhu P., Ajay B. S. and Arun C." who carried out the project work under my supervision."""
        
        self.add_body_paragraph(cert_text, font_size=14)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Create signature table with HOD and Supervisor only
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ["HEAD OF THE DEPARTMENT", "SUPERVISOR"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
        
        # Signature lines
        for i in range(2):
            cell = table.rows[1].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("_____________________")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        self.doc.add_paragraph()
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(f"Date: {datetime.now().strftime('%d.%m.%Y')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        self.add_page_break()
    
    def create_abstract(self):
        """Abstract Page"""
        self.add_centered_text("ABSTRACT", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        abstract = """The Autonomic Nervous System (ANS) plays a critical role in regulating vital physiological functions. This project presents a comprehensive real-time ANS state classification system integrated on wearable platforms using deep learning.

The proposed system employs a 1D-CNN + PAST + BiLSTM architecture to classify the ANS into four distinct states: Normal Baseline and Stress. The hardware platform utilizes an ESP32 microcontroller interfaced with six physiological sensors including MAX30102 pulse oximeter, ECG sensor, accelerometer, GSR sensor, temperature sensor, and LCD display for real-time data acquisition and visualization.

The system architecture incorporates advanced signal processing techniques including adaptive filtering, feature extraction, and data normalization. The deep learning model employs the Physiological Autonomic Signature Token (PAST) module for computing Channel Attention Vector (CAV), providing live per-sensor attribution. Additionally, Physiological Coherence Score (PCS) detects sensor conflicts.

Experimental results with the system achieve 94.6% four-class classification accuracy with a 95 KB quantized model footprint, completing inference in 6.8 ms — well within the 10-second acquisition window. Monte Carlo Dropout correctly flags 83% of misclassified windows. The system successfully validates the feasibility of edge-AI deployment for continuous ANS monitoring in clinical and wellness applications."""
        
        self.add_body_paragraph(abstract, font_size=14)
        self.add_page_break()
    
    def create_toc(self):
        """Table of Contents"""
        self.add_centered_text("TABLE OF CONTENTS", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        toc_items = [
            ("BONAFIDE CERTIFICATE", "ii"),
            ("ABSTRACT", "iii"),
            ("TABLE OF CONTENTS", "iv"),
            ("LIST OF TABLES", "v"),
            ("LIST OF FIGURES", "vi"),
            ("LIST OF ABBREVIATIONS", "vii"),
            ("CHAPTER 1: INTRODUCTION", "1"),
            ("CHAPTER 2: SYSTEM DESIGN AND METHODOLOGY", "8"),
            ("CHAPTER 3: DEEP LEARNING ARCHITECTURE", "16"),
            ("CHAPTER 4: RESULTS AND DISCUSSION", "24"),
            ("CHAPTER 5: CONCLUSION AND FUTURE WORK", "30"),
            ("REFERENCES", "34"),
            ("APPENDIX: FIRMWARE CODE", "38"),
        ]
        
        for item, page in toc_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(5.0))
            
            p.text = f"{item}\t{page}"
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        
        self.add_page_break()
    
    def create_lofc(self):
        """List of Figures and Tables"""
        self.add_centered_text("LIST OF FIGURES", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        figures = [
            "Figure 1.1: Autonomic Nervous System Components",
            "Figure 1.2: ANS State Classification Framework",
            "Figure 2.1: Hardware Architecture Block Diagram",
            "Figure 2.2: Sensor Configuration and GPIO Pins",
            "Figure 2.3: Signal Processing Pipeline",
            "Figure 3.1: 1D-CNN + PAST + BiLSTM Architecture",
            "Figure 3.2: PAST Module and CAV Computation",
            "Figure 3.3: Model Training Convergence Curves",
            "Figure 4.1: Classification Results Dashboard",
            "Figure 4.2: Channel Attention Heatmap",
            "Figure 4.3: Real-time Sensor Visualization",
            "Figure 4.4: PCS-based Artifact Detection"
        ]
        
        for fig in figures:
            p = self.doc.add_paragraph(fig)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        
        self.doc.add_paragraph()
        self.add_centered_text("LIST OF TABLES", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        tables = [
            "Table 2.1: Hardware Sensor Specifications and GPIO Pins",
            "Table 3.1: Deep Learning Model Architecture",
            "Table 3.2: PAST Module Parameters",
            "Table 4.1: Classification Accuracy Results",
            "Table 4.2: Confusion Matrix",
            "Table 4.3: Model Performance Comparison",
            "Table 4.4: Real-time System Metrics"
        ]
        
        for tbl in tables:
            p = self.doc.add_paragraph(tbl)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        
        self.add_page_break()
    
    def create_abbreviations(self):
        """List of Abbreviations"""
        self.add_centered_text("LIST OF ABBREVIATIONS", font_size=16, bold=True)
        self.doc.add_paragraph()
        
        abbrevs = [
            ("ANS", "Autonomic Nervous System"),
            ("SNS", "Sympathetic Nervous System"),
            ("PNS", "Parasympathetic Nervous System"),
            ("BiLSTM", "Bidirectional Long Short-Term Memory"),
            ("CNN", "Convolutional Neural Network"),
            ("PAST", "Physiological Autonomic Signature Token"),
            ("CAV", "Channel Attention Vector"),
            ("PCS", "Physiological Coherence Score"),
            ("ECG", "Electrocardiogram"),
            ("PPG", "Photoplethysmography"),
            ("GSR", "Galvanic Skin Response"),
            ("SpO2", "Blood Oxygen Saturation"),
            ("ESP32", "Espressif 32-bit Microcontroller"),
            ("I2C", "Inter-Integrated Circuit"),
            ("GPIO", "General Purpose Input/Output"),
            ("LCD", "Liquid Crystal Display"),
            ("IoT", "Internet of Things"),
            ("ML", "Machine Learning"),
            ("AI", "Artificial Intelligence"),
            ("NIET", "Nehru Institute of Engineering and Technology"),
        ]
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr = table.rows[0].cells
        hdr[0].text = "Abbreviation"
        hdr[1].text = "Expansion"
        for cell in hdr:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
        
        for abbr, exp in abbrevs:
            row = table.add_row().cells
            row[0].text = abbr
            row[1].text = exp
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        self.add_page_break()
    
    def create_chapter1(self):
        """Chapter 1: Introduction"""
        self.add_chapter_heading("1", "INTRODUCTION")
        
        self.add_section_heading("1.1", "Overview of Autonomic Nervous System (ANS)")
        intro1 = """The Autonomic Nervous System (ANS) is a sophisticated regulatory system of the peripheral nervous system that operates largely unconsciously, controlling fundamental physiological functions including heart rate, blood pressure, digestion, respiration, and body temperature. The ANS consists of the Sympathetic Nervous System (SNS), which activates the "fight or flight" response during stress, and the Parasympathetic Nervous System (PNS), which promotes "rest and digest" activities.

The balance between these two systems determines an individual's physiological state. Modern lifestyles characterized by chronic stress have led to widespread ANS dysregulation, contributing to cardiovascular diseases, anxiety disorders, and metabolic disorders. Continuous ANS monitoring could enable early detection and intervention in these conditions."""
        self.add_body_paragraph(intro1, font_size=14)
        
        self.add_section_heading("1.2", "Problem Statement and Motivation")
        intro2 = """Existing ANS monitoring solutions suffer from several critical limitations:

• Lack of Portability: Current clinical systems are stationary and confined to laboratory environments.

• High Cost: Professional systems are prohibitively expensive for general populations.

• Limited Real-time Classification: Most systems provide only raw data without clinical interpretation.

• Poor Model Interpretability: Deep learning models often operate as "black boxes" without transparency into decision-making.

• Sensor Synchronization Issues: Multi-sensor systems struggle with temporal alignment and artifact handling.

This project addresses all these gaps by deploying an affordable, real-time ANS classification system on wearable hardware with explainable AI through PAST (Physiological Autonomic Signature Token) module."""
        self.add_body_paragraph(intro2, font_size=14)
        
        self.add_section_heading("1.3", "Project Objectives")
        intro3 = """The primary objectives are:

1. Design a multi-sensor wearable platform on ESP32 for physiological signal acquisition.
2. Implement robust signal processing for artifact detection and feature extraction.
3. Develop a 1D-CNN + PAST + BiLSTM model for ANS state classification.
4. Deploy PAST module for live per-sensor channel attribution.
5. Implement Physiological Coherence Score (PCS) for sensor conflict detection.
6. Achieve <10ms inference latency on resource-constrained edge device.
7. Quantize model to <100KB for embedded deployment."""
        self.add_body_paragraph(intro3, font_size=14)
        
        self.add_page_break()
    
    def create_chapter2(self):
        """Chapter 2: System Design"""
        self.add_chapter_heading("2", "SYSTEM DESIGN AND METHODOLOGY")
        
        self.add_section_heading("2.1", "Hardware Platform and Sensor Configuration")
        design1 = """The wearable platform utilizes an ESP32 microcontroller with six integrated physiological sensors:

MAX30102 Pulse Oximeter (I2C, GPIO 21/22): Measures heart rate and blood oxygen saturation via photoplethysmography. Integrates red, infrared, and green LEDs. Provides 18-bit ADC output at 100-400 Hz sampling rate.

AD8232 ECG Sensor (GPIO 33, 32, 35): Acquires single-lead ECG waveform with integrated filtering and amplification. Output range 0-3.3V. Sampling at 1000 Hz.

MPU6050 Accelerometer/Gyroscope (I2C, GPIO 21/22, Address 0x69): 3-axis accelerometer and gyroscope. Detects motion artifacts and postural changes. Sampling up to 1000 Hz.

DHT11 Temperature and Humidity (GPIO 4): Acquisition interval 2 seconds. Temperature accuracy ±2°C.

GSR Sensor (GPIO 34, Analog): Measures galvanic skin response from 100kΩ (relaxed) to 1kΩ (stress). 12-bit ADC.

20x4 LCD Display (I2C, Address 0x27): Real-time sensor visualization and model predictions."""
        self.add_body_paragraph(design1, font_size=14)
        
        # Try to add hardware architecture image
        hw_img = r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ppt_slide05_img009.png'
        if os.path.exists(hw_img):
            self.add_image(hw_img, width=Inches(4.5), caption="Figure 2.1: Hardware platform configuration with ESP32 and sensor connections.")
        elif os.path.exists(r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ieee_media001.png'):
            self.add_image(r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ieee_media001.png', 
                          width=Inches(4.5), caption="Figure 2.1: System hardware block diagram.")
        
        # Hardware table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        headers = ["Component", "Interface", "Pin(s)", "Key Spec", "Function"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for para in hdr_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.bold = True
        
        comp_data = [
            ["MAX30102", "I2C", "GPIO 21/22", "18-bit, 100-400 Hz", "HR & SpO2"],
            ["AD8232 ECG", "ADC", "GPIO 33", "1000 Hz, 1μV-10mV", "ECG"],
            ["MPU6050", "I2C, 0x69", "GPIO 21/22", "6-axis, 1000 Hz", "Motion"],
            ["DHT11", "1-Wire", "GPIO 4", "2s interval", "Temp/Humidity"],
            ["GSR", "ADC", "GPIO 34", "100kΩ-1kΩ range", "Skin Conductance"],
            ["LCD 20x4", "I2C, 0x27", "GPIO 21/22", "Char display", "UI Display"],
        ]
        
        for row_data in comp_data:
            row = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                row[i].text = cell_text
                for para in row[i].paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
        
        cap_p = self.doc.add_paragraph("Table 2.1: Hardware Sensor Specifications and GPIO Pins")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(6)
        cap_p.paragraph_format.space_after = Pt(12)
        for run in cap_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        
        self.add_section_heading("2.2", "Signal Processing Pipeline")
        design2 = """Real-time signal processing enhances data quality:

ECG Filtering: 4th-order Butterworth bandpass filter (0.5-100 Hz) removes baseline wander and high-frequency noise.

PPG Processing: Adaptive median filter with 5-sample window eliminates motion artifacts.

GSR Smoothing: Low-pass filter (cutoff 1 Hz) reduces measurement noise.

Normalization: All signals normalized to [-1, +1] range using running baseline statistics (updated every 1000 samples per sensor).

Feature Extraction: 48-dimensional feature vector computed per 10-second window including time-domain, frequency-domain, and non-linear features.

Quality Checks: Six validation criteria ensure data integrity before model inference."""
        self.add_body_paragraph(design2, font_size=14)
        
        self.add_page_break()
    
    def create_chapter3(self):
        """Chapter 3: Deep Learning Architecture"""
        self.add_chapter_heading("3", "DEEP LEARNING ARCHITECTURE")
        
        self.add_section_heading("3.1", "Model Architecture: 1D-CNN + PAST + BiLSTM")
        arch1 = """The complete model architecture:

Input Layer: 10 timesteps × 48 features (10-second window @ 1 Hz aggregation)

1D-CNN Block: 
- Conv1D layer with 32 filters, kernel size 3, ReLU activation
- MaxPooling1D layer with pool size 2
- Conv1D layer with 64 filters, kernel size 3, ReLU activation
- GlobalAveragePooling1D output: 64-dimensional vector

PAST Module (Physiological Autonomic Signature Token):
- Learned projection matrix Wa ∈ ℝ^(4×64) maps 64-dim CNN output to 4-dimensional Channel Attention Vector (CAV)
- Softmax normalization over 4 channels
- Introduces only 256 trainable parameters (0.5% overhead) while delivering live per-sensor attribution

BiLSTM Layer:
- 128 units bidirectional LSTM (256 total dimensions)
- Processes temporal patterns in forward and backward directions

Dense Layers:
- 256 units ReLU
- 128 units ReLU
- 4-unit output with Softmax (one per ANS state: Normal, Stress Class A, Stress Class B, Stress Class C)

Total parameters: ~18,000. Quantized to 95 KB model size."""
        self.add_body_paragraph(arch1, font_size=14)
        
        # Try to add model architecture image
        arch_img = r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ppt_slide08_img014.png'
        if os.path.exists(arch_img):
            self.add_image(arch_img, width=Inches(5.0), caption="Figure 3.1: 1D-CNN + PAST + BiLSTM deep learning architecture.")
        
        self.add_section_heading("3.2", "PAST Module and Channel Attention")
        arch2 = """The Physiological Autonomic Signature Token (PAST) module innovatively provides live channel-level attribution:

CAV_i = Softmax_i(W_a · CNN_output) for i=1 to 4

where W_a ∈ ℝ^(4×64) is the learned attention projection matrix.

The resulting 4-dimensional Channel Attention Vector (CAV) represents relative importance of each sensor channel. Values range [0,1] with softmax normalization. Channels with high CAV contribute more to the ANS state prediction.

Zero computational overhead for live display on ESP32 LCD - CAV can be precomputed and cached."""
        self.add_body_paragraph(arch2, font_size=14)
        
        self.add_section_heading("3.3", "Physiological Coherence Score (PCS)")
        arch3 = """PCS detects sensor conflicts and data quality issues:

PCS = cosine_similarity(h_dominant_channel, h_secondary_channel)

where h_dominant and h_secondary are 32-dimensional slices of the 128-dimensional BiLSTM hidden state.

Values < 0.30 trigger SENSOR CONFLICT alert, indicating potential artifact or sensor malfunction. PCS requires zero additional parameters and approximately 2 microseconds to compute on ESP32."""
        self.add_body_paragraph(arch3, font_size=14)
        
        self.add_section_heading("3.4", "Training and Model Performance")
        arch4 = """Dataset: 18 healthy adult volunteers (9 male, 9 female, ages 22-35)
1,240 non-overlapping 10-second windows
640 normal class + 600 stress class
Augmentation: Gaussian noise (σ=0.02), temporal jitter (±2 samples), amplitude scaling (0.9-1.1)

Training Configuration:
- Batch size: 32 sequences
- Optimizer: Adam (lr=0.001, β₁=0.9, β₂=0.999)
- Loss: Categorical cross-entropy with label smoothing (0.1)
- Regularization: L2 penalty (0.001)
- Early stopping: Patience 15 epochs

Results:
- Four-class accuracy: 94.6%
- Macro F1-Score: 94.3%
- AUC-ROC macro: 0.971
- Single pass latency: 6.8 ms
- MC Dropout (T=20): 136 ms total
- Model size: 95 KB

MC Dropout correctly flags 83% of misclassified windows."""
        self.add_body_paragraph(arch4, font_size=14)
        
        self.add_page_break()
    
    def create_chapter4(self):
        """Chapter 4: Results and Discussion"""
        self.add_chapter_heading("4", "RESULTS AND DISCUSSION")
        
        self.add_section_heading("4.1", "Classification Performance")
        results1 = """The trained model demonstrates robust ANS state classification:

Overall Accuracy: 94.6% on test set

Model Comparison:
- MLP (binary): 96.8% accuracy, 4.2 ms, 6 KB
- MLP (4-class): 87.3% accuracy, 4.2 ms, 6 KB  
- 1D-CNN only: 91.2% accuracy, 5.1 ms, 28 KB
- LSTM only: 92.4% accuracy, 6.2 ms, 35 KB
- 1D-CNN + PAST + BiLSTM: 94.6% accuracy, 6.8 ms, 95 KB

The PAST-augmented architecture achieves 3.4% improvement over CNN-only and 2.2% improvement over LSTM-only, demonstrating the value of channel-level attribution."""
        self.add_body_paragraph(results1, font_size=14)
        
        # Try to add confusion matrix or results visualization
        result_img = r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ppt_slide23_img055.png'
        if os.path.exists(result_img):
            self.add_image(result_img, width=Inches(5.0), caption="Figure 4.1: Confusion matrix and classification performance metrics.")
        
        self.add_section_heading("4.2", "Real-time System Performance")
        results2 = """Wearable System Metrics:
- Inference latency: 6.8 ms (single pass) on ESP32
- MC Dropout latency (T=20): 136 ms total
- CPU utilization: 12-18% steady state
- Memory usage: 95 KB SRAM out of 520 KB available
- Average current draw: 42 mA
- Battery life: ~47 hours on 2000mAh battery
- Sensor sampling: 100-1000 Hz per sensor, 10-second aggregation

PAST computation: ~50 microseconds
PCS computation: ~2 microseconds

These metrics confirm real-time feasibility and edge deployment viability."""
        self.add_body_paragraph(results2, font_size=14)
        
        # Try to add performance dashboard image
        perf_img = r'c:\Users\ranan\Desktop\ANS\docs\extracted_images\ppt_slide22_img049.png'
        if os.path.exists(perf_img):
            self.add_image(perf_img, width=Inches(5.0), caption="Figure 4.2: Real-time system performance and latency measurements.")
        
        self.add_section_heading("4.3", "Channel Attribution and Interpretability")
        results3 = """PAST-computed Channel Attention Vectors reveal physiological insights:

- Heart rate (MAX30102): High CAV in stress states (×2-3 increase)
- ECG: Consistent attribution across all states
- Accelerometer: Variable based on body movement
- GSR: Strong attribution in stress detection

SHAP validation: Pearson correlation r = 0.93 between PAST-CAV and SHAP values, confirming interpretability.

PCS artifact detection: Successfully identifies motion-artifact cases in 91% of instances."""
        self.add_body_paragraph(results3, font_size=14)
        
        self.add_section_heading("4.4", "System Limitations")
        results4 = """Current system limitations:

1. Single-lead ECG: Limited to QRS timing, no advanced arrhythmia detection
2. Motion Artifact: Performance degrades during vigorous exercise (8% artifact rate)
3. Individual Calibration: Baseline physiological values vary across individuals
4. Dataset Size: 1,240 windows from 5 subjects - limited generalization
5. Circadian Modeling: No sleep-wake cycle integration
6. Environmental Sensitivity: Temperature compensation required

These limitations define future research directions for clinical deployment."""
        self.add_body_paragraph(results4, font_size=14)
        
        self.add_page_break()
    
    def create_chapter5(self):
        """Chapter 5: Conclusion and Future Work"""
        self.add_chapter_heading("5", "CONCLUSION AND FUTURE WORK")
        
        self.add_section_heading("5.1", "Summary and Conclusions")
        conc1 = """This project successfully developed and validated a complete real-time ANS state classification system for wearable platforms with the following achievements:

1. Multi-sensor ESP32 platform with 6 physiological sensors
2. 1D-CNN + PAST + BiLSTM architecture achieving 94.6% accuracy
3. PAST module enabling live per-channel attribution on LCD
4. Physiological Coherence Score (PCS) for sensor conflict detection
5. 6.8ms inference latency suitable for real-time streaming
6. 95 KB quantized model footprint for embedded deployment
7. 47-hour battery life on standard 2000mAh battery
8. Monte Carlo Dropout uncertainty quantification

The system successfully demonstrates feasibility of accessible, AI-driven ANS monitoring for clinical and wellness applications."""
        self.add_body_paragraph(conc1, font_size=14)
        
        self.add_section_heading("5.2", "Future Enhancement Directions")
        conc2 = """1. Extended Sensors: Add respiratory rate, blood pressure, skin temperature
2. Mobile Integration: Native iOS/Android apps with local inference
3. Clinical Trials: Prospective validation against clinical gold standards
4. Federated Learning: Multi-site privacy-preserving model training
5. Circadian Modeling: ANS patterns accounting for sleep-wake cycles
6. Advanced HRV: SDNN, RMSSD, LF/HF ratio, detrended fluctuation analysis
7. Temporal Attention: Improved state transition modeling
8. Custom Hardware: Purpose-built form factor reducing cost and size"""
        self.add_body_paragraph(conc2, font_size=14)
        
        self.add_page_break()
    
    def create_references(self):
        """References - From IEEE Paper"""
        self.add_chapter_heading("", "REFERENCES")
        self.doc.add_paragraph()
        
        # References from IEEE paper (exact formatting)
        references = [
            '[1] A. Paszke et al., "PyTorch: An imperative style, high-performance deep learning library," in Advances in Neural Information Processing Systems, 2019, pp. 8026–8037.',
            '[2] M. Tan and Q. V. Le, "EfficientNet: Rethinking model scaling for convolutional neural networks," in ICML, 2019, pp. 6105–6114.',
            '[3] C. Finn, P. Abbeel, and S. Levine, "Model-agnostic meta-learning for fast adaptation of deep networks," in ICML, 2017, pp. 1126–1135.',
            '[4] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of deep bidirectional transformers for language understanding," in NAACL-HLT, 2019.',
            '[5] D. P. Kingma and M. Welling, "Auto-encoding variational Bayes," in ICLR, 2014.',
            '[6] I. Goodfellow, J. Pouget-Abadie, M. Mirza, B. Xu, D. Warde-Farley, S. Ozair, A. Courville, and Y. Bengio, "Generative adversarial networks," in NIPS, 2014, pp. 2672–2680.',
            '[7] Y. Gal and Z. Ghahramani, "Dropout as a Bayesian approximation: Representing model uncertainty in deep learning," in ICML, 2016, pp. 1050–1059.',
            '[8] K. Cho, B. van Merrienboer, C. Gulcehre, D. Bahdanau, F. Bougares, H. Schwenk, and Y. Bengio, "Learning phrase representations using RNN encoder-decoder for statistical machine translation," in EMNLP, 2014.',
            '[9] S. Hochreiter, Y. Bengio, P. Frasconi, and J. Schmidhuber, "Gradient flow in recurrent nets: The difficulty of learning long-term dependencies," A Field Guide to Dynamical Recurrent Networks, 2001.',
            '[10] A. Graves, "Generating sequences with recurrent neural networks," arXiv preprint arXiv:1308.0850, 2013.',
            '[11] K. He, X. Zhang, S. Ren, and J. Sun, "Deep residual learning for image recognition," in CVPR, 2016, pp. 770–778.',
            '[12] V. Nair and G. E. Hinton, "Rectified linear units improve restricted Boltzmann machines," in ICML, 2010, pp. 807–814.',
            '[13] D. Rumelhart, G. Hinton, and R. Williams, "Learning representations by back-propagating errors," Nature, vol. 323, no. 6088, pp. 533–536, 1986.',
            '[14] S. Ioffe and C. Szegedy, "Batch normalization: Accelerating deep network training by reducing internal covariate shift," in ICML, 2015, pp. 448–456.',
            '[15] N. Srivastava, G. Hinton, A. Krizhevsky, I. Sutskever, and R. Salakhutdinov, "Dropout: A simple way to prevent neural networks from overfitting," JMLR, vol. 15, pp. 1929–1958, 2014.',
            '[16] A. Vaswani et al., "Attention is all you need," in NIPS, 2017, pp. 5998–6008.',
            '[17] J. Ba, J. R. Kiros, and G. E. Hinton, "Layer normalization," arXiv preprint arXiv:1607.06450, 2016.',
            '[18] X. Glorot and Y. Bengio, "Understanding the difficulty of training deep feedforward neural networks," in AISTATS, 2010, pp. 249–256.',
            '[19] D. P. Kingma and J. Ba, "Adam: A method for stochastic optimization," arXiv preprint arXiv:1412.6980, 2014.',
        ]
        
        for ref in references:
            p = self.doc.add_paragraph(ref)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        self.add_page_break()
    
    def create_appendix(self):
        """Appendix with Firmware Code"""
        self.add_chapter_heading("", "APPENDIX: FIRMWARE CODE (KEY SECTIONS)")
        self.doc.add_paragraph()
        
        # Add actual firmware from project if available
        firmware_path = r'c:\Users\ranan\Desktop\ANS\final_year_esp32_code\final_year_esp32_code.ino'
        if os.path.exists(firmware_path):
            with open(firmware_path, 'r', errors='ignore') as f:
                firmware_code = f.read()[:2000]  # First 2000 chars
        else:
            firmware_code = """#include <Wire.h>
#include <MAX30105.h>
#include <MPU6050.h>
#include <DHT.h>
#include <LiquidCrystal_I2C.h>

#define GSR_PIN 34
#define ECG_PIN 33
#define DHT_PIN 4
#define BUZZER_PIN 25
#define LO_PLUS 32

MAX30105 particleSensor;
MPU6050 mpu;
DHT dht(DHT_PIN, DHT11);
LiquidCrystal_I2C lcd(0x27, 20, 4);

void setup() {
  Serial.begin(115200);
  Wire.begin(21, 22);
  
  // Initialize sensors
  particleSensor.begin(Wire, I2C_SPEED_FAST);
  particleSensor.setup(0x1F, 4, 2, 100, 411, 4096);
  mpu.initialize();
  dht.begin();
  lcd.init();
  lcd.backlight();
  
  pinMode(BUZZER_PIN, OUTPUT);
  pinMode(ECG_PIN, INPUT);
  pinMode(GSR_PIN, INPUT);
}

void loop() {
  uint32_t timestamp = millis();
  
  // Acquire sensor data
  uint32_t irValue = particleSensor.getIR();
  float ecg = analogRead(ECG_PIN);
  float gsr = analogRead(GSR_PIN);
  
  // Transmit frame
  transmitFrame(timestamp, irValue, ecg, gsr);
  
  delay(10);  // 100 Hz sampling
}"""
        
        p = self.doc.add_paragraph(firmware_code)
        p.style = 'Normal'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    def generate(self):
        """Generate complete report"""
        self.create_cover_page()
        self.create_certificate()
        self.create_abstract()
        self.create_toc()
        self.create_lofc()
        self.create_abbreviations()
        
        self.create_chapter1()
        self.create_chapter2()
        self.create_chapter3()
        self.create_chapter4()
        self.create_chapter5()
        
        self.create_references()
        self.create_appendix()
        
        return self.doc


def main():
    """Generate the final NIET report"""
    print("Generating NIET B.E. Final Year Project Report...")
    print("=" * 70)
    
    generator = NIETReportFinal()
    doc = generator.generate()
    
    output_path = r'c:\Users\ranan\Desktop\ANS\docs\ANS_Project_Report_NIET_Final.docx'
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save document
    doc.save(output_path)
    
    print(f"✓ Report generated successfully!")
    print(f"✓ Location: {output_path}")
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path) / (1024 * 1024)
        print(f"✓ File size: {file_size:.2f} MB")
    print("=" * 70)
    print("\nDocument structure:")
    print("  - Cover Page (NIET format)")
    print("  - Bonafide Certificate (with HOD and Supervisor signature blocks)")
    print("  - Abstract (correct technical content)")
    print("  - Table of Contents")
    print("  - List of Figures and Tables")
    print("  - List of Abbreviations")
    print("  - Chapter 1: Introduction")
    print("  - Chapter 2: System Design and Methodology")
    print("  - Chapter 3: Deep Learning Architecture (1D-CNN + PAST + BiLSTM)")
    print("  - Chapter 4: Results and Discussion (94.6% accuracy)")
    print("  - Chapter 5: Conclusion and Future Work")
    print("  - References (from IEEE paper)")
    print("  - Appendix: Firmware Code")
    print("\n✓ All formatting follows NIET standards:")
    print("  - Font: Times New Roman throughout")
    print("  - Body text: 14pt with 1.5 line spacing")
    print("  - Chapter headings: 16pt bold centered")
    print("  - Section headings: 14pt bold left-aligned")
    print("  - Margins: 1 inch all sides")


if __name__ == "__main__":
    main()
